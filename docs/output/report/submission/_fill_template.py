# -*- coding: utf-8 -*-
"""Fill official prelim template → 01-作品说明文档翻斗花园.docx"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[4]
TPL = next((ROOT / "assets" / "backup" / "official-2026").glob("*.docx"))
OUT_DIR = Path(__file__).resolve().parent
OUT_DOCX = OUT_DIR / "01-作品说明文档翻斗花园.docx"

IMG_POSTER = ROOT / "assets" / "poster" / "poster-phase1.png"
IMG_ARCH = ROOT / "assets" / "images" / "readme" / "architecture-phase1.png"
IMG_SEQ = ROOT / "assets" / "images" / "readme" / "sequence-phase1.png"


def set_run_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text, size=10.5, bold=False, space_after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading_cn(doc, text, size=16):
    return add_para(doc, text, size=size, bold=True, space_after=10)


def replace_para_text(para, new_text, size=12, bold=False):
    if para.runs:
        para.runs[0].text = new_text
        set_run_font(para.runs[0], size=size, bold=bold)
        for r in para.runs[1:]:
            r.text = ""
    else:
        run = para.add_run(new_text)
        set_run_font(run, size=size, bold=bold)


def set_cell_text(cell, text, size=9):
    for i, p in enumerate(cell.paragraphs):
        if i == 0:
            if p.runs:
                p.runs[0].text = text
                set_run_font(p.runs[0], size=size)
                for r in p.runs[1:]:
                    r.text = ""
            else:
                run = p.add_run(text)
                set_run_font(run, size=size)
        else:
            for r in p.runs:
                r.text = ""


def unique_cells(row):
    seen = []
    out = []
    for c in row.cells:
        tid = id(c._tc)
        if tid not in seen:
            seen.append(tid)
            out.append(c)
    return out


def fill_member_row(row, m):
    cells = row.cells
    mapping_idx = [
        (0, m["name"]),
        (1, m["school"]),
        (2, m["college"]),
        (3, m["major"]),
        (5, m["grade"]),
        (6, m["grad"]),
        (8, m["phone"]),
        (10, m["email"]),
        (11, m["role"]),
    ]
    for idx, val in mapping_idx:
        if idx < len(cells):
            set_cell_text(cells[idx], val)


def main():
    doc = Document(str(TPL))

    mapping = {
        "参赛学校：_______________________________": "参赛学校：中北大学",
        "团队名称：_______________________________": "团队名称：翻斗花园",
        "作品名称：_______________________________": "作品名称：LUT-SA",
        "赛题方向：_______________________________": "赛题方向：操作系统智能创新",
        "联系人（队长）：__________________________": "联系人（队长）：聂君奋",
        "联系电话（队长）：________________________": "联系电话（队长）：【待队长填写】",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in mapping:
            replace_para_text(p, mapping[t], size=12)
        if "请填写作品名称" in t:
            replace_para_text(p, "《 LUT-SA 》 作品原创性声明", size=14, bold=True)
        if t == "一句话抽述关键创新点，不超过30字":
            replace_para_text(p, "将LUT低比特推理做成鸿蒙系统加速服务", size=10.5)
        if t.startswith("提供创意效果图"):
            replace_para_text(
                p,
                "本作品属操作系统智能赛题，以下给出系统架构与调用时序技术方案图（见后文附图），"
                "并附宣传海报。边界：用户态 SystemAbility 封装，不修改 OS 内核。",
                size=10.5,
            )
        if t.startswith("具体描述创意的内容和实现路径"):
            replace_para_text(
                p,
                "见后文「三、介绍文档」正文（不超过800字）及「四、测试报告」。",
                size=10.5,
            )

    table = doc.tables[0]
    set_cell_text(table.rows[0].cells[1], "LUT-SA")
    set_cell_text(table.rows[1].cells[1], "翻斗花园")
    set_cell_text(table.rows[2].cells[1], "中北大学（以队长学校为准）")
    set_cell_text(
        table.rows[3].cells[1],
        "（ ）应用创新　（ ）Agent创新　（ ）用户体验创新　（√）操作系统智能创新",
    )

    members = [
        {
            "name": "聂君奋",
            "school": "中北大学",
            "college": "【待填院系】",
            "major": "【待填专业】",
            "grade": "【待填】",
            "grad": "【待填】",
            "phone": "【待填】",
            "email": "【待填】",
            "role": "队长",
        },
        {
            "name": "范腾达",
            "school": "中北大学",
            "college": "【待填院系】",
            "major": "【待填专业】",
            "grade": "【待填】",
            "grad": "【待填】",
            "phone": "【待填】",
            "email": "【待填】",
            "role": "队员",
        },
        {
            "name": "郑李惠杰",
            "school": "中北大学",
            "college": "【待填院系】",
            "major": "【待填专业】",
            "grade": "【待填】",
            "grad": "【待填】",
            "phone": "【待填】",
            "email": "【待填】",
            "role": "队员",
        },
    ]
    for i, m in enumerate(members):
        fill_member_row(table.rows[6 + i], m)

    for row in table.rows:
        for cell in row.cells:
            if "可列举描述团队" in cell.text or "成员个人或集体" in cell.text:
                set_cell_text(
                    cell,
                    "团队具备系统软件与 AI Agent 工程基础：熟悉文档驱动开发与协作工作流；"
                    "队长统筹方案与系统封装路径，队员分工文档/评测与工程实现。"
                    "指导教师信息由队长补充后定稿。",
                )

    for row in table.rows:
        c0 = row.cells[0].text.strip()
        if c0.startswith("陈") or (c0.endswith("**") and "南京" not in "".join(c.text for c in row.cells)):
            # skip student samples already overwritten; advisor sample often after member rows
            pass

    # advisor: scan for 职称 header then next data row, or remaining 陈**
    advisor_header_idx = None
    for ri, row in enumerate(table.rows):
        joined = "".join(c.text for c in row.cells)
        if "团队指导教师信息" in joined:
            advisor_header_idx = ri
        if advisor_header_idx is not None and ri == advisor_header_idx + 2:
            uniq = unique_cells(row)
            vals = [
                "【待填指导教师姓名】",
                "【待填院系】",
                "【待填职称】",
                "【待填研究方向】",
                "【待填电话】",
                "【待填邮箱】",
            ]
            for u, v in zip(uniq, vals):
                set_cell_text(u, v)
            break

    # Append content
    doc.add_page_break()
    add_heading_cn(doc, "一、创意描述", size=16)
    add_para(doc, "将LUT低比特推理做成鸿蒙系统加速服务", size=10.5)
    add_para(doc, "（共19字，符合不超过30字要求）", size=9)

    add_heading_cn(doc, "二、技术方案（操作系统智能）", size=16)
    add_para(
        doc,
        "整体路径：T-MAC LUT Kernel（计算核心）→ NDK Native 库 → 用户态 SystemAbility（LUT-SA）→ "
        "轻量感知与调度策略 → 上层 AI 任务 / 演示应用。对照基线为 llama.cpp 反量化 CPU 路径。",
        size=10.5,
    )
    add_para(doc, "图1 系统海报", size=10.5, bold=True)
    if IMG_POSTER.exists():
        doc.add_picture(str(IMG_POSTER), width=Cm(14))
    add_para(
        doc,
        "图2 系统架构（L1 应用 → L2 感知调度 → L3 SystemAbility → L4 LUT Kernel）",
        size=10.5,
        bold=True,
    )
    if IMG_ARCH.exists():
        doc.add_picture(str(IMG_ARCH), width=Cm(14))
    add_para(
        doc,
        "图3 调用时序（Demo App → SAMgr → LUT SA → Native Lib → LUT Kernel）",
        size=10.5,
        bold=True,
    )
    if IMG_SEQ.exists():
        doc.add_picture(str(IMG_SEQ), width=Cm(14))

    add_heading_cn(doc, "三、介绍文档", size=16)
    for t in [
        "低比特大模型在边缘设备落地时，主流路径仍以反量化后高精度乘加实现，反量化访存与转换开销显著抵消低比特收益；多任务并发场景下，首字延迟（TTFT）与能耗指标进一步劣化。赛题要求在现有操作系统架构下提升系统AI任务运行效率，纯应用层Demo难以承载「系统级创新」这一命题。",
        "作品LUT-SA以开源T-MAC（EuroSys 2025）比特级查找表范式为计算引擎，将低比特混合精度矩阵乘转换为查表与加法运算，消除反量化乘加路径。计算内核经NDK封装后，以用户态SystemAbility注册到OpenHarmony/HarmonyOS系统服务框架，对外统一暴露模型加载、推理与指标回传三类能力。调度侧联动系统通知、前台状态与负载信号，动态调整AI任务优先级与预取策略，形成「感知→调度→加速」闭环。系统路径全程运行于用户态，不依赖未公开的内核接口。",
        "工程实现分三层：计算层适配T-MAC LUT Kernel至ARM架构，处理访存模式与LUT表布局；服务层按SAMgr注册LUT SystemAbility，对外暴露统一加速接口，应用按需获取proxy；策略层将行为信号映射为优先级、预取与节流策略，可在演示场景中直观呈现。",
        "公开评测数据表明，相对llama.cpp反量化基线，T-MAC在多种边缘CPU上吞吐具备稳定优势。BitNet-3B在M2-Ultra单核22.08对6.49 tokens/s，树莓派5单核8.03对1.37 tokens/s。本作品以此为性能基线，并规划在OpenHarmony模拟器或ARM设备上完成复现与实测补充。复赛阶段可进一步叠加官方NPU/CANN量化路径作对照，强化「异构调度」的论证链。",
        "作品意义在于将LUT计算范式产品化为系统级服务，降低端侧大模型部署门槛，为隐私本地推理、低功耗生成等场景提供基础支撑。",
    ]:
        add_para(doc, t, size=10.5)

    add_heading_cn(doc, "四、测试报告", size=16)
    add_para(doc, "4.1 测试目的", size=12, bold=True)
    add_para(
        doc,
        "验证比特级查找表（LUT）相对反量化基线在边缘CPU上的吞吐优势，并说明封装为OpenHarmony/HarmonyOS用户态系统服务后的复现与移植路径，支撑「系统级AI任务效率提升」主张。",
        size=10.5,
    )
    add_para(doc, "4.2 测试对象与基线", size=12, bold=True)
    add_para(
        doc,
        "加速方案：T-MAC LUT Kernel（本仓二开，上游microsoft/T-MAC）。对比基线：llama.cpp CPU反量化路径。"
        "模型示例：BitNet-3B；Llama-2-7B（W2/W4）。数据来源：公开profiling（docs/profiling_data.md），非本队现场实测。",
        size=10.5,
    )
    add_para(doc, "4.3 公开对比结果（单位：tokens/s）", size=12, bold=True)

    t = doc.add_table(rows=7, cols=6)
    t.style = "Table Grid"
    headers = ["模型", "设备", "线程", "llama.cpp", "T-MAC", "约倍速"]
    data = [
        ["BitNet-3B", "M2-Ultra", "1", "6.49", "22.08", "~3.4×"],
        ["BitNet-3B", "M2-Ultra", "4", "22.09", "54.46", "~2.5×"],
        ["BitNet-3B", "树莓派5", "1", "1.37", "8.03", "~5.9×"],
        ["BitNet-3B", "树莓派5", "2", "2.71", "11.09", "~4.1×"],
        ["Llama-2-7B(W2)", "M2-Ultra", "1", "3.82", "16.68", "~4.4×"],
        ["Llama-2-7B(W2)", "AGX Orin", "1", "0.79", "4.36", "~5.5×"],
    ]
    for j, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[j], h, size=9)
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            set_cell_text(t.rows[i + 1].cells[j], val, size=9)

    add_para(doc, "", size=6)
    add_para(doc, "4.4 复现与移植计划", size=12, bold=True)
    add_para(
        doc,
        "复现：克隆本仓或上游T-MAC，按README与docs/e2e.md准备依赖，转换/获取低比特模型后运行benchmark，并与同配置llama.cpp对照。",
        size=10.5,
    )
    add_para(
        doc,
        "移植计划：①抽取LUT Kernel为Native库；②DevEco打通最小推理；③封装SystemAbility（Load/Infer/Metrics）；"
        "④叠加感知调度并对照；⑤固化本平台TTFT、tokens/s及可选功耗数据。",
        size=10.5,
    )
    add_para(doc, "4.5 风险与边界", size=12, bold=True)
    add_para(
        doc,
        "初赛数据为公开profiling，非本队板测，文中已标注来源；ARM鸿蒙设备存在ISA/内存布局差异，优先用户态库；"
        "SystemAbility能力级别以可申请、可演示为准。",
        size=10.5,
    )

    add_heading_cn(doc, "五、待队长补全清单（提交前）", size=16)
    for item in [
        "封面与信息表：队长电话、三名队员院系/专业/年级/毕业时间/电话/邮箱",
        "指导教师：姓名、院系、职称、研究方向、电话、邮箱（须与队长同校）并完成审核签名",
        "原创性声明页：全体队员签名 + 指导老师签名（扫描或电子签）",
        "导出PDF，命名：01-作品说明文档翻斗花园.pdf",
    ]:
        add_para(doc, "· " + item, size=10.5)

    doc.save(str(OUT_DOCX))
    print("saved", OUT_DOCX)
    print("size", OUT_DOCX.stat().st_size)


if __name__ == "__main__":
    main()
